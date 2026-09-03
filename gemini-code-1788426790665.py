import base64
import io
import os
import json
import tempfile
import urllib.request
from pathlib import Path
from datetime import date, datetime
import uuid
import sqlite3

import streamlit as st

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ADMIN_PASSWORD = "admin123"

st.set_page_config(
    page_title="Sang loc Thalassemia",
    page_icon="🩸",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={"Get Help": None, "Report a bug": None},
)

CSS = """<style>
.stApp { background: linear-gradient(180deg, #E6F4F8 0%, #F4F8FB 100%); }
.block-container { padding-top: 1.2rem; padding-bottom: 3.5rem; max-width: 1180px; }
header, footer, .viewerBadge_container__1QSob, div[data-testid="stDecoration"] { display: none; }
.hero { background: linear-gradient(135deg, #012A4A 0%, #014F86 50%, #0077B6 100%);
        color: #fff; border-radius: 22px; padding: 26px 30px; margin-bottom: 16px; }
.hero h1 { color: #fff !important; font-size: 30px !important; margin: 0 0 8px 0 !important; }
.hero p { color: #D9F3FF; font-size: 16px; margin: 0; line-height: 1.6; }
h2 { color: #012A4A !important; }
h3 { color: #014F86 !important; }
.stButton > button { background: linear-gradient(180deg, #0077B6, #014F86) !important;
    color: #fff !important; border: 0 !important; border-radius: 12px !important;
    font-weight: 700 !important; min-height: 3.3rem !important; }
.info-box, .warning-box, .danger-box, .success-box {
    padding: 14px 16px; border-radius: 14px; margin: 10px 0; line-height: 1.7; }
.info-box { background: #E7F6FB; border-left: 6px solid #0077B6; }
.warning-box { background: #FFF6E8; border-left: 6px solid #F4A261; }
.danger-box { background: #FDECEC; border-left: 6px solid #D62828; }
.success-box { background: #E9F7EF; border-left: 6px solid #2A9D8F; }

/* CSS Thẻ bệnh viện 3 dòng đóng khung */
.hospital-box { 
    background-color: #FFFFFF; 
    border: 1px solid #BEE3F8; 
    border-left: 5px solid #0077B6; 
    border-radius: 10px; 
    padding: 12px 14px; 
    margin-bottom: 12px; 
    box-shadow: 0 2px 4px rgba(0,0,0,0.04);
}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

def box(kind, text):
    st.markdown('<div class="%s">%s</div>' % (kind, text), unsafe_allow_html=True)

VUNG_MIEN = [
    "Chọn vùng/miền", "Đông Bắc", "Tây Bắc", "Đồng bằng sông Hồng", "Bắc Trung Bộ",
    "Trung Trung Bộ", "Nam Trung Bộ", "Tây Nguyên", "Đông Nam Bộ", "Đồng bằng sông Cửu Long"
]

TINH_THEO_VUNG = {
    "Đông Bắc": ["Hà Giang", "Cao Bằng", "Bắc Kạn - Thái Nguyên", "Tuyên Quang", "Lạng Sơn", "Bắc Giang"],
    "Tây Bắc": ["Điện Biên - Lai Châu", "Sơn La", "Hòa Bình", "Yên Bái"],
    "Đồng bằng sông Hồng": ["Hà Nội", "Vĩnh Phúc - Phú Thọ", "Bắc Ninh - Hưng Yên", "Quảng Ninh", "Hải Dương", "Thái Bình", "Hải Phòng", "Nam Định - Ninh Bình", "Thanh Hóa"],
    "Bắc Trung Bộ": ["Nghệ An", "Hà Tĩnh", "Quảng Bình", "Quảng Trị - Thừa Thiên Huế"],
    "Trung Trung Bộ": ["Đà Nẵng", "Quảng Nam - Quảng Ngãi", "Bình Định", "Phú Yên"],
    "Nam Trung Bộ": ["Khánh Hòa", "Ninh Thuận - Bình Thuận"],
    "Tây Nguyên": ["Kon Tum", "Gia Lai", "Đắc Lắc", "Đắc Nông - Lâm Đồng"],
    "Đông Nam Bộ": ["Bình Phước - Bình Dương", "Đồng Nai - Bà Rịa Vũng Tàu", "Tây Ninh - Long An", "Thành phố Hồ Chí Minh"],
    "Đồng bằng sông Cửu Long": ["Tiền Giang - Vĩnh Long", "Bến Tre - Trà Vinh", "Đồng Tháp - An Giang", "Kiên Giang - Hậu Giang", "Cần Thơ", "Sóc Trăng - Bạc Liêu - Cà Mau"],
}

TAT_CA_TINH = ["Chọn Tỉnh/Thành phố"]
for ds in TINH_THEO_VUNG.values():
    TAT_CA_TINH.extend(ds)
VUNG_CUA_TINH = {tinh: vung for vung, ds in TINH_THEO_VUNG.items() for tinh in ds}

DAN_TOC = [
    "Chọn dân tộc", "Stiêng", "Ê Đê", "Gia Rai", "Ba Na", "Xơ Đăng", "Cơ Ho", "Hrê",
    "Chăm", "Khơ Me", "Thái", "Mường", "Tày", "Nùng", "Dao", "Sán Chay", "Kinh", "Hoa", "Dân tộc khác"
]
DIEM_DAN_TOC_VN = {
    "Stiêng": 3.0, "Ê Đê": 3.0, "Gia Rai": 3.0, "Ba Na": 2.5, "Xơ Đăng": 2.5, "Cơ Ho": 2.5, "Hrê": 2.5,
    "Chăm": 2.0, "Khơ Me": 2.0, "Thái": 2.0, "Mường": 2.0, "Tày": 1.5, "Nùng": 1.5,
    "Dao": 1.5, "Sán Chay": 1.5, "Kinh": 0.5, "Hoa": 0.5, "Dân tộc khác": 0.5,
}

NGHIEP_NHOM = [
    "Chọn nghề nghiệp", "Nông nghiệp - Lâm nghiệp - Thủy sản", "Lao động phổ thông / Công nhân",
    "Học sinh - Sinh viên", "Cán bộ - Công chức - Viên chức", "Kinh doanh - Dịch vụ - Văn phòng",
    "Y tế - Giáo dục - Kỹ thuật", "Nghỉ hưu / Nội trợ", "Khác"
]

ALTITUDE_OPTIONS = [
    "Dưới 1.000m (Đồng bằng / Trung duy / Núi thấp)",
    "1.000m – 1.499m (-0.8 g/dL)",
    "1.500m – 1.999m (-1.1 g/dL)",
    "2.000m – 2.499m (-1.4 g/dL)",
    "2.500m – 2.999m (-1.8 g/dL)",
    "3.000m – 3.499m (-2.1 g/dL)",
    "3.500m – 3.999m (-2.5 g/dL)",
    "4.000m – 4.499m (-2.9 g/dL)",
    "Từ 4.500m trở lên (-3.3 g/dL)"
]

ALTITUDE_CORRECTION_MAP = {
    "Dưới 1.000m (Đồng bằng / Trung duy / Núi thấp)": 0.0,
    "1.000m – 1.499m (-0.8 g/dL)": 0.8,
    "1.500m – 1.999m (-1.1 g/dL)": 1.1,
    "2.000m – 2.499m (-1.4 g/dL)": 1.4,
    "2.500m – 2.999m (-1.8 g/dL)": 1.8,
    "3.000m – 3.499m (-2.1 g/dL)": 2.1,
    "3.500m – 3.999m (-2.5 g/dL)": 2.5,
    "4.000m – 4.499m (-2.9 g/dL)": 2.9,
    "Từ 4.500m trở lên (-3.3 g/dL)": 3.3
}

CAU_HOI = [
    "1. Bản thân hoặc người thân trong gia đình đã từng được chẩn đoán mắc bệnh Thalassemia / Thiếu máu truyền máu chưa?",
    "2. Gia đình/dòng họ có ai bị biến dạng xương mặt, lách to, da vàng hoặc da sạm đen bất thường không?",
    "3. Bản thân đã từng có tiền sử xét nghiệm ghi nhận thiếu máu nhược sắc, hồng cầu nhỏ hoặc nghi ngờ mang gen chưa?",
    "4. Vợ/Chồng hoặc người chuẩn bị kết hôn có thuộc cùng dòng họ, cùng dân tộc ít người hoặc cùng thôn/bản không?",
    "5. Vợ/Chồng hoặc bạn đời đã từng được xét nghiệm sàng lọc Thalassemia chưa?",
    "6. Đã từng bị từ chối hiến máu do nồng độ Huyết sắc tố (Hb) thấp hoặc hồng cầu quá nhỏ chưa?",
    "7. Bản thân thường xuyên có biểu hiện mệt mỏi, hoa mắt, chóng mặt, da xanh xao kéo dài không?",
    "8. Đã từng thực hiện xét nghiệm Điện di Huyết sắc tố (Hb electrophoresis) hoặc xét nghiệm Gen Thalassemia chưa?",
    "9. Kết quả xét nghiệm trước đây có ghi nhận biến thể Hb (như HbE, HbCS) hoặc mang gen dị hợp tử không?",
    "10. Đã từng được bác sĩ chuyên khoa huyết học tư vấn về nguy cơ sinh con mắc bệnh Thalassemia thể nặng chưa?",
    "11. Trong tiền sử gia đình có ghi nhận trường hợp sảy thai liên tiếp, thai chết lưu không rõ nguyên nhân hoặc phù thai không?",
    "12. Nơi sinh sống hiện tại hoặc quê quán thuộc khu vực miền núi, vùng sâu vùng xa có tỷ lệ mang gen cao không?",
    "13. Bố và Mẹ đẻ có phải là người cùng một dân tộc thiểu số không?",
    "14. Bố hoặc Mẹ đẻ đã từng được xác định mang gen bệnh Thalassemia chưa?",
    "15. Bạn có nhu cầu tư vấn di truyền trước hôn nhân hoặc trước khi sinh con không?",
]

def bv(ten, diachi, dt, hang=None):
    return {"ten": ten, "diachi": diachi, "dt": dt, "hang": hang}

BENH_VIEN = {
    "Hà Giang": [bv("Bệnh viện Đa khoa Tỉnh Hà Giang", "185 Trần Hưng Đạo, TP. Hà Giang", "0219 3851 215", 2)],
    "Cao Bằng": [bv("Bệnh viện Đa khoa Tỉnh Cao Bằng", "06 Phan Đình Phùng, TP. Cao Bằng", "0219 3852 317", 2)],
    "Bắc Kạn - Thái Nguyên": [
        bv("Bệnh viện Đa khoa Tỉnh Bắc Kạn", "01 Nguyễn Thị Minh Khai, TP. Bắc Kạn", "0209 3812 245", 2),
        bv("Bệnh viện Đa khoa Trung ương Thái Nguyên", "479 Lương Ngọc Quyến, TP. Thái Nguyên", "0208 3852 345", 1)
    ],
    "Tuyên Quang": [bv("Bệnh viện Đa khoa Tỉnh Tuyên Quang", "Đường Lê Duẩn, TP. Tuyên Quang", "0207 3822 512", 2)],
    "Lạng Sơn": [bv("Bệnh viện Đa khoa Tỉnh Lạng Sơn", "Thôn Nhị Hà, Xã Hoàng Đồng, TP. Lạng Sơn", "0205 3812 280", 2)],
    "Bắc Giang": [bv("Bệnh viện Đa khoa Tỉnh Bắc Giang", "Đường Lê Lợi, TP. Bắc Giang", "0204 3854 289", 2)],
    "Điện Biên - Lai Châu": [bv("Bệnh viện Đa khoa Tỉnh Điện Biên", "Phường Mường Thanh, TP. Điện Biên Phủ", "0215 3825 211", 2)],
    "Sơn La": [bv("Bệnh viện Đa khoa Tỉnh Sơn La", "Tổ 4, Phường Chiềng Sinh, TP. Sơn La", "0212 3852 232", 2)],
    "Hòa Bình": [bv("Bệnh viện Đa khoa Tỉnh Hòa Bình", "Phường Đồng Tiến, TP. Hòa Bình", "0218 3852 018", 2)],
    "Yên Bái": [bv("Bệnh viện Đa khoa Tỉnh Yên Bái", "Thôn 1, Xã Tiền Phong, TP. Yên Bái", "0216 3852 240", 2)],
    "Hà Nội": [
        bv("Bệnh viện Bạch Mai (Khoa Huyết học)", "78 Giải Phóng, Đống Đa, Hà Nội", "024 3869 3731", 1),
        bv("Viện Huyết học - Truyền máu Trung ương", "Phố Phạm Văn Bạch, Cầu Giấy, Hà Nội", "024 3782 1895", 1)
    ],
    "Vĩnh Phúc - Phú Thọ": [bv("Bệnh viện Đa khoa Tỉnh Phú Thọ", "Đường Nguyễn Tất Thành, TP. Việt Trì", "0210 6254 115", 1)],
    "Bắc Ninh - Hưng Yên": [bv("Bệnh viện Đa khoa Tỉnh Bắc Ninh", "Đường Nguyễn Quyền, TP. Bắc Ninh", "0222 3821 242", 2)],
    "Quảng Ninh": [bv("Bệnh viện Đa khoa Tỉnh Quảng Ninh", "Phố Tuệ Tĩnh, P. Hồng Hải, TP. Hạ Long", "0203 3825 478", 1)],
    "Hải Dương": [bv("Bệnh viện Đa khoa Tỉnh Hải Dương", "224 Nguyễn Lương Bằng, TP. Hải Dương", "0220 3890 205", 2)],
    "Thái Bình": [bv("Bệnh viện Đa khoa Tỉnh Thái Bình", "530 Lý Bôn, TP. Thái Bình", "0227 3831 031", 1)],
    "Hải Phòng": [bv("Bệnh viện Hữu nghị Việt Tiệp", "1 Nhà Thương, Lê Chân, Hải Phòng", "0225 3700 436", 1)],
    "Nam Định - Ninh Bình": [bv("Bệnh viện Đa khoa Tỉnh Nam Định", "02 Trần Quốc Toản, TP. Nam Định", "0228 3849 233", 2)],
    "Thanh Hóa": [bv("Bệnh viện Đa khoa Tỉnh Thanh Hóa", "181 Hải Thượng Lãn Ông, TP. Thanh Hóa", "0237 3951 042", 1)],
    "Nghệ An": [bv("Bệnh viện Hữu nghị Đa khoa Nghệ An", "Đại lộ Lê Nin, TP. Vinh", "0238 3844 528", 1)],
    "Hà Tĩnh": [bv("Bệnh viện Đa khoa Tỉnh Hà Tĩnh", "01 Hải Thượng Lãn Ông, TP. Hà Tĩnh", "0239 3855 561", 2)],
    "Quảng Bình": [bv("Bệnh viện Hữu nghị Việt Nam - Cu Ba Đồng Hới", "TK 10, P. Nam Lý, Đồng Hới", "0232 3822 216", 1)],
    "Quảng Trị - Thừa Thiên Huế": [bv("Bệnh viện Trung ương Huế", "16 Lê Lợi, TP. Huế", "0234 3822 325", 1)],
    "Đà Nẵng": [bv("Bệnh viện Đà Nẵng", "124 Hải Phòng, Q. Hải Châu, Đà Nẵng", "0236 3821 118", 1)],
    "Quảng Nam - Quảng Ngãi": [bv("Bệnh viện Đa khoa Tỉnh Quảng Nam", "14 Lý Thường Kiệt, TP. Tam Kỳ", "0235 3851 523", 2)],
    "Bình Định": [bv("Bệnh viện Đa khoa Tỉnh Bình Định", "106 Nguyễn Huệ, TP. Quy Nhơn", "0256 3822 211", 1)],
    "Phú Yên": [bv("Bệnh viện Đa khoa Tỉnh Phú Yên", "15 Nguyễn Hữu Thọ, TP. Tuy Hòa", "0257 3823 219", 2)],
    "Khánh Hòa": [bv("Bệnh viện Đa khoa Tỉnh Khánh Hòa", "19 Yersin, TP. Nha Trang", "0258 3822 112", 1)],
    "Ninh Thuận - Bình Thuận": [bv("Bệnh viện Đa khoa Tỉnh Bình Thuận", "Trường Chinh, TP. Phan Thiết", "0252 3822 211", 2)],
    "Kon Tum": [bv("Bệnh viện Đa khoa Tỉnh Kon Tum", "224 Bà Triệu, TP. Kon Tum", "0260 3862 573", 2)],
    "Gia Lai": [bv("Bệnh viện Đa khoa Tỉnh Gia Lai", "132 Tôn Thất Tùng, TP. Pleiku", "0269 3824 402", 2)],
    "Đắc Lắc": [bv("Bệnh viện Đa khoa Vùng Tây Nguyên", "184 Trần Quý Cáp, TP. Buôn Ma Thuột", "0262 3852 234", 1)],
    "Đắc Nông - Lâm Đồng": [bv("Bệnh viện Đa khoa Tỉnh Lâm Đồng", "01 Phạm Ngọc Thạch, TP. Đà Lạt", "0263 3822 115", 2)],
    "Bình Phước - Bình Dương": [bv("Bệnh viện Đa khoa Tỉnh Bình Dương", "512 Phạm Ngọc Thạch, TP. Thủ Dầu Một", "0274 3822 153", 1)],
    "Đồng Nai - Bà Rịa Vũng Tàu": [bv("Bệnh viện Đa khoa Đồng Nai", "2 Đồng Khởi, TP. Biên Hòa", "0251 8969 999", 1)],
    "Tây Ninh - Long An": [bv("Bệnh viện Đa khoa Tỉnh Long An", "211 Nguyễn Thông, TP. Tân An", "0272 3826 330", 2)],
    "Thành phố Hồ Chí Minh": [
        bv("Bệnh viện Chợ Rẫy", "201 Nguyễn Chí Thanh, Quận 5, TP.HCM", "028 3855 4137", 1),
        bv("Bệnh viện Truyền máu Huyết học TP.HCM", "1 Trần Hữu Trang, Q. Tân Bình, TP.HCM", "028 3839 7535", 1)
    ],
    "Tiền Giang - Vĩnh Long": [bv("Bệnh viện Đa khoa Tỉnh Tiền Giang", "Ấp 3, Xã Trung An, TP. Mỹ Tho", "0273 3872 363", 2)],
    "Bến Tre - Trà Vinh": [bv("Bệnh viện Nguyễn Đình Chiểu", "Đoàn Hoàng Minh, TP. Bến Tre", "0275 3822 532", 2)],
    "Đồng Tháp - An Giang": [bv("Bệnh viện Đa khoa Trung tâm An Giang", "60 Hải Lãn Ông, TP. Long Xuyên", "0296 3852 543", 1)],
    "Kiên Giang - Hậu Giang": [bv("Bệnh viện Đa khoa Tỉnh Kiên Giang", "13 A Lâm Quang Ky, TP. Rạch Giá", "0297 3862 044", 1)],
    "Cần Thơ": [bv("Bệnh viện Đa khoa Trung ương Cần Thơ", "315 Nguyễn Văn Linh, Q. Ninh Kiều, Cần Thơ", "0292 3820 071", 1)],
    "Sóc Trăng - Bạc Liêu - Cà Mau": [bv("Bệnh viện Đa khoa Tỉnh Cà Mau", "16 tháng 4, Phường 6, TP. Cà Mau", "0290 3831 015", 2)]
}

def lay_benh_vien_theo_tinh(tinh_tru):
    if not tinh_tru or tinh_tru == "Chọn Tỉnh/Thành phố":
        return [], "khong", ""
    if tinh_tru in BENH_VIEN:
        return BENH_VIEN[tinh_tru], "chinh", ""
    vung_tru = VUNG_CUA_TINH.get(tinh_tru, "")
    ds = []
    for tinh, lst in BENH_VIEN.items():
        if VUNG_CUA_TINH.get(tinh, "") == vung_tru:
            ds.extend(lst)
    return ds, ("goi_y" if ds else "khong"), vung_tru

def classify_round1(s1):
    if s1 <= 5:
        return ("Nguy cơ thấp", "Theo dõi sức khỏe định kỳ. Chưa cần làm công thức máu ngay.")
    if s1 <= 12:
        return ("Nguy cơ trung bình", "Khuyến nghị làm công thức máu (CBC) tại cơ sở y tế gần nhất.")
    return "Nguy cơ cao", "Cần làm công thức máu và thực hiện xét nghiệm chuyên sâu Vòng 2 & 3."

def classify_total(s1, s2):
    total = s1 + s2
    if total >= 15:
        return ("Nguy cơ RẤT CAO", "Gửi mẫu điện di huyết sắc tố / Xét nghiệm gen ngay. Tư vấn di truyền cho cả hai vợ chồng.")
    if total >= 10 or s2 >= 4:
        return ("Nguy cơ CAO", "Thực hiện điện di huyết sắc tố. Tư vấn di truyền chi tiết.")
    if total >= 6:
        return ("Nguy cơ TRUNG BÌNH", "Theo dõi, kiểm tra công thức máu lại sau 6 tháng.")
    return "Nguy cơ THẤP", "Tư vấn sức khỏe sinh sản / theo dõi định kỳ."

def score_round2(mcv, mch, hb_tho, rbc, rdw, gioitinh, do_cao_option):
    giam_hb = ALTITUDE_CORRECTION_MAP.get(do_cao_option, 0.0)
    hb_hieuchinh = hb_tho - giam_hb  
    
    hb_cut = 12.0 if gioitinh == "Nữ" else 13.0
    
    score_mcv = 4 if (mcv > 0 and mcv < 86.0) else 0
    score_mch = 3 if (mch > 0 and mch < 27.0) else 0
    score_hb = 2 if (hb_hieuchinh > 0 and hb_hieuchinh < hb_cut) else 0
    score_rbc = 2 if rbc >= 5.0 else 0
    score_rdw = 2 if (11.0 <= rdw <= 15.0 and mcv > 0 and mcv < 86.0) else 0

    str_hb_display = f"{hb_hieuchinh:.2f} g/dL (đã trừ {giam_hb:.1f} g/dL)" if giam_hb > 0 else f"{hb_hieuchinh:.2f} g/dL"

    rows = [
        ("MCV (Thể tích hồng cầu)", f"{mcv:.2f} fL", "86.0 – 98.0 fL", score_mcv),
        ("MCH (Hb trung bình HC)", f"{mch:.2f} pg", "27.0 – 32.0 pg", score_mch),
        ("Hemoglobin (Hb hiệu chỉnh)", str_hb_display, f"< {hb_cut:.1f} g/dL ({gioitinh})", score_hb),
        ("RBC (Số lượng hồng cầu)", f"{rbc:.2f} M/uL", "4.30 – 5.80 M/uL", score_rbc),
        ("RDW (Độ phân bố HC)", f"{rdw:.2f} %", "11.0 – 15.0 %", score_rdw),
    ]
    return sum(r[3] for r in rows), rows, hb_hieuchinh

DATA_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "thalassemia_ho_so.db")

def _db_connect():
    conn = sqlite3.connect(DATA_FILE)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS ho_so (
            HoSoID TEXT PRIMARY KEY,
            data TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
    """)
    conn.commit()
    return conn

def serialize_data(obj):
    if isinstance(obj, (date, datetime)):
        return obj.strftime("%d/%m/%Y")
    if isinstance(obj, uuid.UUID):
        return str(obj)
    raise TypeError(f"Type {type(obj)} not serializable")

def doc_du_lieu_luu_tru():
    try:
        conn = _db_connect()
        rows = conn.execute("SELECT data FROM ho_so ORDER BY updated_at DESC").fetchall()
        conn.close()
        return [json.loads(row[0]) for row in rows]
    except Exception:
        return []

def ghi_du_lieu_duy_nhat(record_data):
    try:
        conn = _db_connect()
        cursor = conn.cursor()
        now = datetime.now().isoformat(timespec="seconds")
        
        ho_so_id = str(record_data.get("HoSoID") or uuid.uuid4().hex)
        record_data["HoSoID"] = ho_so_id

        json_str = json.dumps(record_data, ensure_ascii=False, default=serialize_data)

        cursor.execute(
            "INSERT OR REPLACE INTO ho_so (HoSoID, data, updated_at) VALUES (?, ?, ?)",
            (ho_so_id, json_str, now),
        )
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Lỗi ghi dữ liệu: {e}")
        return False

def init_state():
    defaults = {
        "page": "home", "is_admin": False, "danh_sach_khach_hang": [], "ho_so_id": "",
        "s1": 0.0, "s2": 0, "r2_detail": None, "go_r2": False, "go_r3": False,
        "gioitinh": "Nữ", "hoten": "", "ngaysinh": date(2000, 1, 1), "dantoc": "Chọn dân tộc",
        "vung_o": "Chọn vùng/miền", "vung_lamviec": "Chọn vùng/miền", "tinh_o": "Chọn Tỉnh/Thành phố",
        "tinh_lamviec": "Chọn Tỉnh/Thành phố", "nghenghiep": "Chọn nghề nghiệp", "sdt": "", "diachi": "",
        "do_cao": ALTITUDE_OPTIONS[0],
        "kieu_hb": "Chưa rõ", "hb_bienthe": "", "hba": 0.0, "hba2": 0.0, "hbf": 0.0,
        "ketluan_r3": "", "tuvan_r3": "", "ghichu": "", "mcv": 0.0, "mch": 0.0, "hb": 0.0,
        "rbc": 0.0, "rdw": 0.0, "vong1_da_luu": False, "vong2_da_luu": False, "vong3_da_luu": False, "lan_luu_cuoi": "",
    }
    for i in range(1, 16): 
        defaults["q%s" % i] = "Không"
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def luu_thong_tin_khach_hang(vong=None):
    ss = st.session_state
    ho_ten = str(ss.get("hoten", "")).strip()
    if not ho_ten:
        st.warning("Vui lòng nhập Họ tên trước khi lưu!")
        return False

    if not ss.get("ho_so_id"): ss.ho_so_id = uuid.uuid4().hex

    if vong == 1: ss.vong1_da_luu = True
    elif vong == 2: ss.vong2_da_luu = True
    elif vong == 3: ss.vong3_da_luu = True

    ngay_sinh_val = ss.get("ngaysinh")
    if isinstance(ngay_sinh_val, (date, datetime)):
        ngay_sinh_val = ngay_sinh_val.strftime("%d/%m/%Y")

    cau_tra_loi_v1 = []
    for i in range(1, 16):
        cau_tra_loi_v1.append({
            "stt": i,
            "cau_hoi": CAU_HOI[i-1],
            "tra_loi": ss.get(f"q{i}", "Không")
        })

    data = {
        "HoSoID": ss.ho_so_id,
        "ThoiGian": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "HoTen": ho_ten,
        "GioiTinh": ss.get("gioitinh", "Nữ"),
        "NgaySinh": ngay_sinh_val,
        "DanToc": ss.get("dantoc", ""),
        "SDT": str(ss.get("sdt", "")).strip(),
        "DiaChi": str(ss.get("diachi", "")).strip(),
        "TinhO": ss.get("tinh_o", ""),
        "VungO": ss.get("vung_o", ""),
        "TinhLamViec": ss.get("tinh_lamviec", ""),
        "VungLamViec": ss.get("vung_lamviec", ""),
        "NgheNghiep": ss.get("nghenghiep", ""),
        "DoCaoSinhSong": ss.get("do_cao", ALTITUDE_OPTIONS[0]),
        
        "DiemV1": ss.get("s1", 0.0) if ss.vong1_da_luu else None,
        "KetLuanV1": classify_round1(ss.get("s1", 0.0))[0] if ss.vong1_da_luu else "Chưa lưu V1",
        "DeXuatV1": classify_round1(ss.get("s1", 0.0))[1] if ss.vong1_da_luu else "",
        "CauTraLoiV1": cau_tra_loi_v1,
        
        "DiemV2": ss.get("s2", 0) if ss.vong2_da_luu else None,
        "ChiTietV2": ss.get("r2_detail") if ss.vong2_da_luu else None,
        "KetLuanTong": classify_total(ss.get("s1", 0.0), ss.get("s2", 0))[0] if ss.vong2_da_luu else "Chưa lưu V2",
        "DeXuatTong": classify_total(ss.get("s1", 0.0), ss.get("s2", 0))[1] if ss.vong2_da_luu else "",
        "MCV": ss.get("mcv", 0.0) if ss.vong2_da_luu else None,
        "MCH": ss.get("mch", 0.0) if ss.vong2_da_luu else None,
        "Hb": ss.get("hb", 0.0) if ss.vong2_da_luu else None,
        "RBC": ss.get("rbc", 0.0) if ss.vong2_da_luu else None,
        "RDW": ss.get("rdw", 0.0) if ss.vong2_da_luu else None,
        
        "HbA": ss.get("hba", 0.0) if ss.vong3_da_luu else None,
        "HbA2": ss.get("hba2", 0.0) if ss.vong3_da_luu else None,
        "HbF": ss.get("hbf", 0.0) if ss.vong3_da_luu else None,
        "KieuHb": ss.get("kieu_hb", "") if ss.vong3_da_luu else "",
        "HbBienThe": ss.get("hb_bienthe", "") if ss.vong3_da_luu else "",
        "KetLuanV3": ss.get("ketluan_r3", "") if ss.vong3_da_luu else "Chưa lưu V3",
        "TuVanV3": ss.get("tuvan_r3", "") if ss.vong3_da_luu else "",
        "GhiChu": ss.get("ghichu", "") if ss.vong3_da_luu else "",
        
        "Vong1DaLuu": ss.vong1_da_luu,
        "Vong2DaLuu": ss.vong2_da_luu,
        "Vong3DaLuu": ss.vong3_da_luu,
    }

    ok = ghi_du_lieu_duy_nhat(data)
    if ok:
        ss.danh_sach_khach_hang = doc_du_lieu_luu_tru()
        ss.lan_luu_cuoi = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    return ok

def tao_phieu_word_tu_data(data_dict):
    ds_bv_o, _, _ = lay_benh_vien_theo_tinh(data_dict.get("TinhO", ""))
    ds_bv_lam, _, _ = lay_benh_vien_theo_tinh(data_dict.get("TinhLamViec", ""))
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(1.5)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10)

    def set_cell_shading(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def set_cell_text(cell, text, bold=False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        r = p.add_run(str(text))
        r.bold = bold
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        r.font.size = Pt(9.5)

    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            set_cell_text(cell, h, True)
            set_cell_shading(cell, "014F86")
            p = cell.paragraphs[0]
            for run in p.runs:
                rPr = run._element.get_or_add_rPr()
                color = OxmlElement('w:color')
                color.set(qn('w:val'), 'FFFFFF')
                rPr.append(color)

        for row_idx, row in enumerate(rows):
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                cell = row_cells[i]
                set_cell_text(cell, value)
                set_cell_shading(cell, "F3FAFD" if row_idx % 2 == 1 else "FFFFFF")
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        return table

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        r.font.size = Pt(11)
        rPr = r._element.get_or_add_rPr()
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '014F86')
        rPr.append(color)

    for text, size, bold in [
        ("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", 12, True),
        ("Độc lập - Tự do - Hạnh phúc", 10, True),
        ("PHIẾU KẾT QUẢ SÀNG LỌC THALASSEMIA", 14, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        r.font.size = Pt(size)
        if size == 14 and bold:
            rPr = r._element.get_or_add_rPr()
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '014F86')
            rPr.append(color)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.add_run(f"Thời gian lập phiếu: {data_dict.get('ThoiGian', '')}").font.size = Pt(8.5)

    add_heading("1. Thông tin cá nhân & Địa bàn sinh sống/làm việc")
    add_table(["Thông tin", "Giá trị khai báo"], [
        ["Họ và tên", data_dict.get("HoTen", "Chưa nhập")],
        ["Ngày sinh", str(data_dict.get("NgaySinh", ""))],
        ["Giới tính", data_dict.get("GioiTinh", "")],
        ["Dân tộc", data_dict.get("DanToc", "")],
        ["Nơi sinh sống", f"{data_dict.get('TinhO', '')} ({data_dict.get('VungO', '')})"],
        ["Nơi làm việc / Học tập", f"{data_dict.get('TinhLamViec', '')} ({data_dict.get('VungLamViec', '')})"],
        ["Độ cao sinh sống (WHO 2024)", data_dict.get("DoCaoSinhSong", ALTITUDE_OPTIONS[0])],
        ["Số điện thoại", data_dict.get("SDT", "Chưa nhập")],
    ])

    add_heading("2. Kết quả Vòng 1 - Tiền sử & Dịch tễ")
    if not data_dict.get("Vong1DaLuu"):
        doc.add_paragraph("Chưa thực hiện Vòng 1.")
    else:
        diem_v1 = data_dict.get("DiemV1")
        add_table(["Chỉ số", "Kết quả"], [
            ["Điểm Vòng 1", "%.1f điểm" % (diem_v1 if diem_v1 is not None else 0)],
            ["Đánh giá nguy cơ", data_dict.get("KetLuanV1", "")],
            ["Khuyến nghị", data_dict.get("DeXuatV1", "")],
        ])

        ds_cautraloi = data_dict.get("CauTraLoiV1", [])
        if ds_cautraloi:
            add_heading("2.1 Chi tiết 15 câu hỏi khảo sát Vòng 1")
            rows_q = [[str(item.get("stt", idx+1)), item.get("cau_hoi", ""), item.get("tra_loi", "Không")] for idx, item in enumerate(ds_cautraloi)]
            add_table(["STT", "Nội dung câu hỏi khảo sát", "Trả lời"], rows_q)

    add_heading("3. Kết quả Vòng 2 - Huyết học (Tham chiếu Phiếu Xét Nghiệm)")
    if not data_dict.get("Vong2DaLuu") or not data_dict.get("ChiTietV2"):
        doc.add_paragraph("Chưa thực hiện Vòng 2.")
    else:
        rows = [[ten, giatri, nguong, str(diem)] for ten, giatri, nguong, diem in data_dict.get("ChiTietV2")]
        diem_v2 = data_dict.get("DiemV2") or 0
        diem_v1 = data_dict.get("DiemV1") or 0
        rows.append(["TỔNG ĐIỂM VÒNG 2", "%d/13" % diem_v2, "Tổng V1+V2", "%.1f" % (diem_v1 + diem_v2)])
        add_table(["Thông số", "Giá trị", "Chỉ số bình thường (CSBT)", "Điểm"], rows)

    add_heading("4. Cơ sở Y tế gần nhất gợi ý khám")
    ds_tong = ds_bv_o + [b for b in ds_bv_lam if b not in ds_bv_o]
    if ds_tong:
        add_table(["Bệnh viện / Cơ sở Y tế", "Địa chỉ", "Điện thoại"], [[x["ten"], x["diachi"], x["dt"]] for x in ds_tong[:3]])
    else:
        doc.add_paragraph("Chưa khai báo địa bàn sinh sống/làm việc.")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def render_main():
    init_state()
    ss = st.session_state

    st.markdown("""
    <div class="hero">
        <h1>🩸 SÀNG LỌC VÀ ĐÁNH GIÁ NGUY CƠ THALASSEMIA</h1>
        <p>Giao diện hỗ trợ tối ưu cho người lớn tuổi • Tự động tính Hb hiệu chỉnh • Chuẩn tham chiếu Phiếu Xét Nghiệm</p>
    </div>
    """, unsafe_allow_html=True)

    with st.sidebar:
        st.title("⚙️ Hệ thống")
        if not ss.is_admin:
            pwd = st.text_input("Mật khẩu cán bộ:", type="password")
            if st.button("Đăng nhập Admin"):
                if pwd == ADMIN_PASSWORD:
                    ss.is_admin = True
                    st.rerun()
                else: st.error("Sai mật khẩu!")
        else:
            st.success("👨‍⚕️ Cán bộ Y tế")
            if st.button("Đăng xuất Admin"):
                ss.is_admin = False
                st.rerun()

        st.markdown("---")
        if st.button("🔄 Tạo hồ sơ mới"):
            for k in list(st.session_state.keys()): del st.session_state[k]
            st.rerun()

    st.subheader("📋 1. Thông tin cá nhân & Địa bàn sinh sống / Làm việc")
    
    col1, col2 = st.columns(2)
    with col1:
        ss.hoten = st.text_input("Họ và tên người khám (*):", value=ss.hoten)
        c_sub1, c_sub2 = st.columns(2)
        with c_sub1:
            ss.gioitinh = st.selectbox("Giới tính:", ["Nữ", "Nam"], index=0 if ss.gioitinh == "Nữ" else 1)
        with c_sub2:
            ss.ngaysinh = st.date_input(
                "Ngày sinh:",
                value=ss.ngaysinh,
                min_value=date(1900, 1, 1),
                max_value=date.today(),
                format="DD/MM/YYYY"
            )
        
        ss.dantoc = st.selectbox("Dân tộc:", DAN_TOC, index=DAN_TOC.index(ss.dantoc) if ss.dantoc in DAN_TOC else 0)
        ss.nghenghiep = st.selectbox("Nghề nghiệp:", NGHIEP_NHOM, index=NGHIEP_NHOM.index(ss.nghenghiep) if ss.nghenghiep in NGHIEP_NHOM else 0)
        ss.sdt = st.text_input("Số điện thoại liên hệ:", value=ss.sdt)

    with col2:
        st.markdown("**🏡 Nơi sinh sống / Thường trú hiện tại:**")
        c_vo, c_to = st.columns(2)
        with c_vo:
            ss.vung_o = st.selectbox("Vùng sinh sống:", VUNG_MIEN, index=VUNG_MIEN.index(ss.vung_o) if ss.vung_o in VUNG_MIEN else 0, key="sb_vung_o")
        with c_to:
            tinh_ds_o = TAT_CA_TINH if ss.vung_o == "Chọn vùng/miền" else ["Chọn Tỉnh/Thành phố"] + TINH_THEO_VUNG.get(ss.vung_o, [])
            ss.tinh_o = st.selectbox("Tỉnh/Thành phố sinh sống:", tinh_ds_o, index=tinh_ds_o.index(ss.tinh_o) if ss.tinh_o in tinh_ds_o else 0, key="sb_tinh_o")

        st.markdown("**🏢 Nơi làm việc / Học tập:**")
        c_vl, c_tl = st.columns(2)
        with c_vl:
            ss.vung_lamviec = st.selectbox("Vùng làm việc:", VUNG_MIEN, index=VUNG_MIEN.index(ss.vung_lamviec) if ss.vung_lamviec in VUNG_MIEN else 0, key="sb_vung_lam")
        with c_tl:
            tinh_ds_lam = TAT_CA_TINH if ss.vung_lamviec == "Chọn vùng/miền" else ["Chọn Tỉnh/Thành phố"] + TINH_THEO_VUNG.get(ss.vung_lamviec, [])
            ss.tinh_lamviec = st.selectbox("Tỉnh/Thành phố làm việc:", tinh_ds_lam, index=tinh_ds_lam.index(ss.tinh_lamviec) if ss.tinh_lamviec in tinh_ds_lam else 0, key="sb_tinh_lam")

        ss.do_cao = st.selectbox("🏔️ Độ cao nơi sinh sống (Tự động trừ Hb):", ALTITUDE_OPTIONS, index=ALTITUDE_OPTIONS.index(ss.do_cao) if ss.do_cao in ALTITUDE_OPTIONS else 0)

    # --- KHU VỰC HIỂN THỊ CƠ SỞ Y TẾ DẠNG THẺ (2 CỘT SONG SONG) ---
    st.markdown("#### 🏥 Cơ sở Y tế gợi ý khám gần nhất")
    col_bv_o, col_bv_lam = st.columns(2)
    
    with col_bv_o:
        st.markdown(f"**🏠 Nơi sinh sống / Thường trú ({ss.tinh_o}):**")
        ds_o, _, _ = lay_benh_vien_theo_tinh(ss.tinh_o)
        if ds_o:
            for bv_item in ds_o:
                html_card = f"""
                <div class="hospital-box">
                    <div style="font-weight: 700; color: #014F86; font-size: 15px; margin-bottom: 4px;">🏥 {bv_item['ten']}</div>
                    <div style="color: #333333; font-size: 13.5px; margin-bottom: 3px;">📍 <b>Địa chỉ:</b> {bv_item['diachi']}</div>
                    <div style="color: #0077B6; font-size: 13.5px; font-weight: 600;">📞 <b>Hotline:</b> {bv_item['dt']}</div>
                </div>
                """
                st.markdown(html_card, unsafe_allow_html=True)
        else:
            st.caption("Chưa chọn Tỉnh/Thành phố sinh sống.")

    with col_bv_lam:
        st.markdown(f"**🏢 Nơi làm việc / Học tập ({ss.tinh_lamviec}):**")
        ds_lam, _, _ = lay_benh_vien_theo_tinh(ss.tinh_lamviec)
        if ds_lam:
            for bv_item in ds_lam:
                html_card = f"""
                <div class="hospital-box">
                    <div style="font-weight: 700; color: #014F86; font-size: 15px; margin-bottom: 4px;">🏥 {bv_item['ten']}</div>
                    <div style="color: #333333; font-size: 13.5px; margin-bottom: 3px;">📍 <b>Địa chỉ:</b> {bv_item['diachi']}</div>
                    <div style="color: #0077B6; font-size: 13.5px; font-weight: 600;">📞 <b>Hotline:</b> {bv_item['dt']}</div>
                </div>
                """
                st.markdown(html_card, unsafe_allow_html=True)
        else:
            st.caption("Chưa chọn Tỉnh/Thành phố làm việc.")

    st.markdown("---")

    tab1, tab2, tab3 = st.tabs(["VÒNG 1: Tiền sử & Dịch tễ", "VÒNG 2: Công thức máu (Nhập từ phiếu)", "VÒNG 3: Điện di Hb & Gen"])

    with tab1:
        st.markdown("### 📝 Bảng câu hỏi khảo sát nhanh")
        score_v1 = DIEM_DAN_TOC_VN.get(ss.dantoc, 0.5)

        for idx, q_text in enumerate(CAU_HOI, 1):
            key = f"q{idx}"
            col_q, col_a = st.columns([3.5, 1])
            with col_q: st.write(q_text)
            with col_a:
                ans = st.radio(f"q_{idx}", ["Có", "Không"], key=key, horizontal=True, label_visibility="collapsed")
                if ans == "Có": score_v1 += 1.0

        ss.s1 = score_v1
        if st.button("💾 Lưu kết quả Vòng 1", key="btn_v1"):
            if luu_thong_tin_khach_hang(vong=1):
                st.success(f"Đã lưu Vòng 1! Tổng điểm: {ss.s1}")
        
        if ss.vong1_da_luu:
            kl_v1, dx_v1 = classify_round1(ss.s1)
            box("info-box", f"<b>Kết luận Vòng 1:</b> {kl_v1}<br/><b>Khuyến nghị:</b> {dx_v1}")

    with tab2:
        st.markdown("### 🔬 Vòng 2: Nhập số liệu trực tiếp từ Phiếu Xét Nghiệm")
        st.caption("💡 *Chỉ cần nhập các chỉ số ghi trên giấy kết quả, phần mềm sẽ tự đối chiếu khoảng tham chiếu CSBT và tự tính độ cao.*")

        c_mcv, c_mch, c_hb, c_rbc, c_rdw = st.columns(5)
        with c_mcv: ss.mcv = st.number_input("MCV (fL) [CSBT: 86-98]:", value=ss.mcv, min_value=0.0, step=0.1)
        with c_mch: ss.mch = st.number_input("MCH (pg) [CSBT: 27-32]:", value=ss.mch, min_value=0.0, step=0.1)
        with c_hb: ss.hb = st.number_input("Hemoglobin (g/dL) [CSBT: 13-18]:", value=ss.hb, min_value=0.0, step=0.1)
        with c_rbc: ss.rbc = st.number_input("RBC (M/uL) [CSBT: 4.3-5.8]:", value=ss.rbc, min_value=0.0, step=0.01)
        with c_rdw: ss.rdw = st.number_input("RDW (%) [CSBT: 11-15]:", value=ss.rdw, min_value=0.0, step=0.1)

        s2_score, r2_rows, hb_hieuchinh = score_round2(ss.mcv, ss.mch, ss.hb, ss.rbc, ss.rdw, ss.gioitinh, ss.do_cao)
        ss.s2 = s2_score
        ss.r2_detail = r2_rows

        if ss.hb > 0:
            st.markdown(f"⚙️ **Kết quả tự động tính toán:** Nồng độ Hb hiệu chỉnh theo độ cao = **{hb_hieuchinh:.2f} g/dL**")

        if st.button("💾 Lưu kết quả Vòng 2", key="btn_v2"):
            if luu_thong_tin_khach_hang(vong=2):
                st.success(f"Đã lưu Vòng 2! Điểm V2: {ss.s2}")

        if ss.vong2_da_luu:
            kl_tong, dx_tong = classify_total(ss.s1, ss.s2)
            box("danger-box" if "CAO" in kl_tong else "success-box", f"<b>Đánh giá tổng hợp nguy cơ:</b> {kl_tong}<br/><b>Hướng xử trí:</b> {dx_tong}")

    with tab3:
        st.markdown("### 🧬 Vòng 3: Xét nghiệm Điện di Huyết sắc tố / Gen")
        c_hba, c_hba2, c_hbf = st.columns(3)
        with c_hba: ss.hba = st.number_input("HbA (%):", value=ss.hba, min_value=0.0, max_value=100.0, step=0.1)
        with c_hba2: ss.hba2 = st.number_input("HbA2 (%):", value=ss.hba2, min_value=0.0, max_value=100.0, step=0.1)
        with c_hbf: ss.hbf = st.number_input("HbF (%):", value=ss.hbf, min_value=0.0, max_value=100.0, step=0.1)

        ss.kieu_hb = st.text_input("Kiểu Hb:", value=ss.kieu_hb)
        ss.ketluan_r3 = st.text_area("Kết luận Vòng 3:", value=ss.ketluan_r3)

        if st.button("💾 Lưu kết quả Vòng 3", key="btn_v3"):
            if luu_thong_tin_khach_hang(vong=3):
                st.success("Đã lưu Vòng 3 thành công!")

    st.markdown("---")
    try:
        cau_tra_loi_v1 = []
        for i in range(1, 16):
            cau_tra_loi_v1.append({
                "stt": i,
                "cau_hoi": CAU_HOI[i-1],
                "tra_loi": ss.get(f"q{i}", "Không")
            })

        data_download = {
            "HoSoID": str(ss.get("ho_so_id", "")),
            "ThoiGian": datetime.now().strftime("%d/%m/%Y %H:%M"),
            "HoTen": ss.get("hoten", "").strip(),
            "GioiTinh": ss.get("gioitinh", ""),
            "NgaySinh": ss.get("ngaysinh").strftime("%d/%m/%Y") if isinstance(ss.get("ngaysinh"), (date, datetime)) else str(ss.get("ngaysinh")),
            "DanToc": ss.get("dantoc", ""),
            "SDT": ss.get("sdt", "").strip(),
            "TinhO": ss.get("tinh_o", ""),
            "VungO": ss.get("vung_o", ""),
            "TinhLamViec": ss.get("tinh_lamviec", ""),
            "VungLamViec": ss.get("vung_lamviec", ""),
            "DoCaoSinhSong": ss.get("do_cao", ALTITUDE_OPTIONS[0]),
            
            "DiemV1": ss.get("s1", 0.0) if ss.vong1_da_luu else None,
            "KetLuanV1": classify_round1(ss.get("s1", 0.0))[0] if ss.vong1_da_luu else "Chưa lưu V1",
            "DeXuatV1": classify_round1(ss.get("s1", 0.0))[1] if ss.vong1_da_luu else "",
            "CauTraLoiV1": cau_tra_loi_v1,
            
            "DiemV2": ss.get("s2", 0) if ss.vong2_da_luu else None,
            "ChiTietV2": ss.get("r2_detail") if ss.vong2_da_luu else None,
            "KetLuanTong": classify_total(ss.get("s1", 0.0), ss.get("s2", 0))[0] if ss.vong2_da_luu else "Chưa lưu V2",
            
            "Vong1DaLuu": ss.vong1_da_luu,
            "Vong2DaLuu": ss.vong2_da_luu,
            "Vong3DaLuu": ss.vong3_da_luu,
        }

        word_bytes = tao_phieu_word_tu_data(data_download)
        st.download_button(
            label="📄 Tải phiếu kết quả (.docx)",
            data=word_bytes,
            file_name=f"PHIEU_SANG_LOC_{ss.get('hoten','').replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    except Exception as e:
        st.error(f"Lỗi chuẩn bị file tải về: {e}")

if __name__ == "__main__":
    render_main()